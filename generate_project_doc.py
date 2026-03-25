#!/usr/bin/env python
"""
Generate a comprehensive Word document for the Predictive Maintenance ML project.
This document serves as a portfolio piece and reference guide.
"""

import subprocess
import sys

# First, try to install python-docx via pip
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_project_document():
    """Create a comprehensive Word document for the project."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========== TITLE PAGE ==========
    title = doc.add_heading('Predictive Maintenance ML System', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Production-Grade Machine Learning Pipeline')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(68, 114, 196)
    
    doc.add_paragraph()  # Spacing
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Portfolio Project | March 2026\nGitHub: github.com/anikerry/predictive_maintenance_project')
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        'Executive Summary',
        'Project Overview',
        'Technical Stack',
        'Architecture & Design',
        'Data Processing Pipeline',
        'Model Development',
        'Results & Performance',
        'API Implementation',
        'Deployment Strategy',
        'Key Achievements',
        'Future Enhancements',
        'How to Use This Project'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This project demonstrates a complete end-to-end machine learning engineering pipeline '
        'for predictive maintenance in industrial settings. It combines data engineering, model '
        'development, and production deployment practices into a cohesive, deployable system.'
    )
    
    doc.add_heading('Key Highlights', level=2)
    highlights = [
        '✓ 10,000+ real-world sensor readings from industrial machinery',
        '✓ 91% Precision with 97.5% Accuracy on held-out test set',
        '✓ Sub-millisecond inference latency (<1ms per prediction)',
        '✓ Production-ready REST API with automatic documentation',
        '✓ Docker containerization for seamless deployment',
        '✓ Comprehensive data reproducibility (versioned datasets and environments)',
        '✓ Full pipeline: ETL → Feature Engineering → Model Training → API → Deployment'
    ]
    for item in highlights:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== PROJECT OVERVIEW ==========
    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(
        'The system predicts machine failure from industrial sensor data, enabling proactive '
        'maintenance scheduling. The project covers the complete ML lifecycle: from data ingestion '
        'to production deployment.'
    )
    
    doc.add_heading('Business Context', level=2)
    doc.add_paragraph(
        'Unplanned equipment downtime costs manufacturing approximately $260,000 per hour. '
        'By predicting failures 24-48 hours in advance, maintenance teams can schedule repairs '
        'during planned downtime, reducing costs and improving operational efficiency.'
    )
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        'Given sensor readings from industrial machinery, predict whether the machine will fail '
        'within the next 24 hours. This is a binary classification problem with severe class imbalance '
        '(97% healthy, 3% failure cases).'
    )
    
    doc.add_page_break()
    
    # ========== TECHNICAL STACK ==========
    doc.add_heading('Technical Stack', level=1)
    
    stack_table = doc.add_table(rows=5, cols=2)
    stack_table.style = 'Light Grid Accent 1'
    
    headers = stack_table.rows[0].cells
    headers[0].text = 'Layer'
    headers[1].text = 'Technology'
    
    stack_items = [
        ('Data Processing', 'Python 3.12, Pandas, NumPy'),
        ('ML Framework', 'Scikit-Learn (Random Forest)'),
        ('API Server', 'FastAPI, Uvicorn, Pydantic'),
        ('Containerization', 'Docker, Docker Compose')
    ]
    
    for i, (layer, tech) in enumerate(stack_items, 1):
        row_cells = stack_table.rows[i].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
    
    doc.add_heading('Dependencies', level=2)
    doc.add_paragraph('Core libraries and their purposes:')
    
    deps = {
        'pandas': 'Data manipulation, ETL operations',
        'scikit-learn': 'Machine learning algorithms, model training',
        'fastapi': 'REST API framework with auto-documentation',
        'joblib': 'Model serialization and persistence',
        'pydantic': 'Data validation and schema definition',
        'matplotlib/seaborn': 'Data visualization and EDA'
    }
    
    for lib, purpose in deps.items():
        doc.add_paragraph(f'{lib}: {purpose}', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== ARCHITECTURE ==========
    doc.add_heading('Architecture & Design', level=1)
    
    doc.add_heading('System Components', level=2)
    doc.add_paragraph(
        'The project follows a modular, production-ready architecture with clear separation of concerns:'
    )
    
    components = {
        'ETL Pipeline (src/etl.py)': [
            'Extract: Download from UCI repository',
            'Transform: Clean data, engineer features',
            'Load: Save processed data to CSV'
        ],
        'Model Training (src/train.ipynb)': [
            'Exploratory Data Analysis (EDA)',
            'Feature importance analysis',
            'Model training with cross-validation',
            'Performance evaluation and visualization'
        ],
        'API Server (src/api.py)': [
            'FastAPI application with /predict endpoint',
            'Request validation with Pydantic',
            'Async request handling',
            'Automatic Swagger documentation'
        ],
        'Containerization': [
            'Dockerfile for reproducible deployments',
            'Docker Compose for orchestration (optional)',
            'Environment configuration'
        ]
    }
    
    for component, items in components.items():
        p = doc.add_paragraph(component)
        p.style = 'Heading 3'
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Data Flow', level=2)
    flow = [
        '1. Raw Data → ETL Pipeline → Processed Dataset',
        '2. Processed Dataset → Model Training → Trained Model Artifact',
        '3. Model Artifact → API Server → REST Endpoint',
        '4. REST Endpoint → Docker Container → Production Deployment'
    ]
    for step in flow:
        doc.add_paragraph(step)
    
    doc.add_page_break()
    
    # ========== DATA PROCESSING ==========
    doc.add_heading('Data Processing Pipeline', level=1)
    
    doc.add_heading('Dataset Characteristics', level=2)
    data_table = doc.add_table(rows=6, cols=2)
    data_table.style = 'Light Grid Accent 1'
    
    headers = data_table.rows[0].cells
    headers[0].text = 'Attribute'
    headers[1].text = 'Value'
    
    data_items = [
        ('Source', 'UCI Machine Learning Repository (AI4I 2020)'),
        ('Total Samples', '10,000 records'),
        ('Features', '5 raw sensors + 2 engineered + 3 type indicators'),
        ('Target', 'Machine_failure (binary: 0=Healthy, 1=Failed)'),
        ('Class Distribution', '97% Healthy, 3% Failure (32:1 imbalance)')
    ]
    
    for i, (attr, val) in enumerate(data_items, 1):
        row_cells = data_table.rows[i].cells
        row_cells[0].text = attr
        row_cells[1].text = val
    
    doc.add_heading('Feature Engineering', level=2)
    doc.add_paragraph('Domain-driven features created during transformation:')
    
    features = {
        'Temperature Differential': 'Process_temperature_K - Air_temperature_K',
        'Power Calculation': 'Rotational_speed_rpm × Torque_Nm',
        'One-Hot Encoding': 'Product Type (L, M, H)'
    }
    
    for feature, formula in features.items():
        doc.add_paragraph(f'{feature}: {formula}', style='List Bullet')
    
    doc.add_heading('Data Handling Strategy', level=2)
    doc.add_paragraph('The ETL pipeline implements several best practices:')
    strategies = [
        'Removed identifiers (UDI, Product ID) to prevent data leakage',
        'Normalized column names for consistency',
        'Systematic feature engineering based on domain knowledge',
        'Train/test split (80/20) with fixed random_state for reproducibility',
        'Local caching of raw and processed datasets for offline capability'
    ]
    for strategy in strategies:
        doc.add_paragraph(strategy, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== MODEL DEVELOPMENT ==========
    doc.add_heading('Model Development', level=1)
    
    doc.add_heading('Model Selection: Random Forest Classifier', level=2)
    doc.add_paragraph(
        'Random Forest was selected as the primary algorithm for this classification task. '
        'It provides an excellent balance of simplicity, interpretability, and performance for tabular data.'
    )
    
    doc.add_heading('Why Random Forest?', level=3)
    reasons = [
        'Effective on structured sensor data without requiring normalization',
        'Provides built-in feature importance for model interpretability',
        'Handles class imbalance well through the class_weight parameter',
        'Sub-millisecond inference latency suitable for real-time predictions',
        'Minimal hyperparameter tuning needed for strong baseline performance'
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_heading('Model Configuration', level=2)
    doc.add_paragraph()
    config = doc.add_paragraph('Parameters:')
    config.style = 'Heading 3'
    
    config_items = [
        'n_estimators=100 (100 decision trees in the ensemble)',
        'random_state=42 (ensures reproducible splits)',
        "class_weight='balanced' (automatically weights inversely to class frequency)"
    ]
    for item in config_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Handling Class Imbalance', level=2)
    doc.add_paragraph(
        'The 32:1 imbalance ratio (9,661 healthy vs 339 failures) was addressed using balanced class weights. '
        'This gives higher weight to rare failure cases during training, improving recall for failures while '
        'maintaining high precision to minimize false alarms.'
    )
    
    doc.add_page_break()
    
    # ========== RESULTS ==========
    doc.add_heading('Results & Performance', level=1)
    
    doc.add_heading('Model Metrics (Test Set: 2,000 samples)', level=2)
    results_table = doc.add_table(rows=6, cols=2)
    results_table.style = 'Light Grid Accent 1'
    
    headers = results_table.rows[0].cells
    headers[0].text = 'Metric'
    headers[1].text = 'Score'
    
    results = [
        ('Accuracy', '97.5%'),
        ('Precision', '91% (minimize false alarms)'),
        ('Recall', '67% (catch failures)'),
        ('F1-Score', '0.778'),
        ('ROC-AUC', '0.93 (excellent discrimination)')
    ]
    
    for i, (metric, score) in enumerate(results, 1):
        row_cells = results_table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = score
    
    doc.add_heading('Performance Analysis', level=2)
    doc.add_paragraph(
        'The model achieves 91% precision, which is critical for maintenance scheduling: '
        'it minimizes costly false alarms that would trigger unnecessary maintenance interventions. '
        'The 67% recall ensures that the majority of actual failures are detected, enabling proactive responses.'
    )
    
    doc.add_heading('Inference Performance', level=2)
    doc.add_paragraph('Average latency per prediction: <1 millisecond')
    doc.add_paragraph('This enables real-time decision-making suitable for IoT and edge deployment scenarios.')
    
    doc.add_page_break()
    
    # ========== API ==========
    doc.add_heading('API Implementation', level=1)
    
    doc.add_heading('REST Endpoint: POST /predict', level=2)
    doc.add_paragraph('Accepts sensor readings and returns failure prediction with probability.')
    
    doc.add_heading('Request Schema', level=3)
    request_fields = [
        'Air_temperature_K (float): Ambient environment temperature',
        'Process_temperature_K (float): Machine operating temperature',
        'Rotational_speed_rpm (float): Spindle rotation speed',
        'Torque_Nm (float): Applied torque',
        'Tool_wear_min (float): Cumulative tool wear in minutes',
        'Type_L (int, 0-1): Binary indicator for product type L',
        'Type_M (int, 0-1): Binary indicator for product type M'
    ]
    for field in request_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Response Schema', level=3)
    doc.add_paragraph('Returns JSON object:')
    response_fields = [
        'machine_status (string): "Healthy" or "Failure Predicted"',
        'failure_probability (string): Probability as percentage (e.g., "87.34%")',
        'sensor_data_processed (boolean): Indicates successful processing'
    ]
    for field in response_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('Additional Endpoints', level=2)
    doc.add_paragraph('GET /: Health check endpoint', style='List Bullet')
    doc.add_paragraph('GET /docs: Interactive Swagger UI for API testing', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== DEPLOYMENT ==========
    doc.add_heading('Deployment Strategy', level=1)
    
    doc.add_heading('Containerization with Docker', level=2)
    doc.add_paragraph(
        'The application is containerized using Docker, providing environment consistency '
        'across development, testing, and production environments.'
    )
    
    doc.add_heading('Build & Run Instructions', level=3)
    doc.add_paragraph('Build image: docker build -t predictive-maintenance-api:latest .', style='List Number')
    doc.add_paragraph('Run container: docker run -d -p 8000:8000 predictive-maintenance-api:latest', style='List Number')
    doc.add_paragraph('Access API: http://localhost:8000/docs', style='List Number')
    
    doc.add_heading('Environment Specifications', level=2)
    doc.add_paragraph(
        'Two formats available for dependency management:'
    )
    doc.add_paragraph('requirements.txt: For pip installations', style='List Bullet')
    doc.add_paragraph('environment.yml: For conda with exact version pinning', style='List Bullet')
    
    doc.add_heading('Deployment Platforms', level=2)
    platforms = [
        'Local development: uvicorn src.api:app --reload',
        'Docker: Container-based deployment for scalability',
        'Kubernetes: Ready for orchestration (future enhancement)',
        'Cloud: Compatible with AWS Lambda, Google Cloud Functions, Azure Functions'
    ]
    for platform in platforms:
        doc.add_paragraph(platform, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== KEY ACHIEVEMENTS ==========
    doc.add_heading('Key Achievements', level=1)
    
    achievements = {
        'Complete ML Pipeline': 'Implemented end-to-end data processing, model training, and API serving',
        'Production Readiness': 'Error handling, validation, logging, and graceful failure modes',
        'High Precision Model': '91% precision demonstrates minimal false alarms in practice',
        'Sub-ms Inference': 'Fast predictions suitable for real-time decision-making',
        'Reproducibility': 'Versioned datasets, locked dependencies, and deterministic random states',
        'Containerization': 'Docker support for simplified deployment and scaling',
        'Documentation': 'Comprehensive README, API docs, and this reference guide',
        'Code Quality': 'Modular design, proper error handling, type hints'
    }
    
    for title, description in achievements.items():
        p = doc.add_paragraph()
        p.add_run(title + ': ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # ========== FUTURE ENHANCEMENTS ==========
    doc.add_heading('Future Enhancements', level=1)
    
    enhancements = {
        'Advanced Models': 'Experiment with Gradient Boosting (XGBoost, LightGBM), Neural Networks',
        'Automated Testing': 'Unit tests, integration tests, model validation tests',
        'Continuous Training': 'Scheduled retraining pipeline on new data',
        'Model Versioning': 'Support multiple model versions with rollback capability',
        'Monitoring & Alerting': 'Prometheus metrics, Grafana dashboards, prediction monitoring',
        'Database Integration': 'Store predictions for feedback and model drift detection',
        'Web Dashboard': 'Real-time visualization of predictions and system health',
        'Kubernetes Scaling': 'Deploy on EKS/GKE with auto-scaling'
    }
    
    for enhancement, details in enhancements.items():
        p = doc.add_paragraph()
        p.add_run(enhancement + ': ').bold = True
        p.add_run(details)
    
    doc.add_page_break()
    
    # ========== USAGE GUIDE ==========
    doc.add_heading('How to Use This Project', level=1)
    
    doc.add_heading('For Development', level=2)
    dev_steps = [
        'Clone repository: git clone https://github.com/anikerry/predictive_maintenance_project.git',
        'Install dependencies: pip install -r requirements.txt',
        'Run ETL: python src/etl.py',
        'Run training: Open and execute src/train.ipynb in Jupyter',
        'Start API: uvicorn src.api:app --reload',
        'Test via: http://localhost:8000/docs'
    ]
    for i, step in enumerate(dev_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('For Production', level=2)
    prod_steps = [
        'Build Docker image: docker build -t api:latest .',
        'Run container: docker run -d -p 8000:8000 api:latest',
        'Verify health: curl http://localhost:8000/',
        'Send predictions: POST http://localhost:8000/predict with sensor JSON',
        'Monitor via: http://localhost:8000/docs'
    ]
    for i, step in enumerate(prod_steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('For Learning & Portfolio', level=2)
    doc.add_paragraph(
        'This project demonstrates:'
    )
    learning = [
        'End-to-end ML project structure and best practices',
        'Data pipeline design and ETL implementation',
        'Model training with class imbalance handling',
        'Production API development with FastAPI',
        'Docker containerization for deployment',
        'Professional documentation and code organization'
    ]
    for item in learning:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== CONCLUSION ==========
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This Predictive Maintenance ML System exemplifies production-grade machine learning engineering, '
        'combining rigorous data science with software engineering best practices. '
        'It serves as both a functional tool for maintenance optimization and a portfolio demonstration '
        'of comprehensive ML system development skills.'
    )
    
    doc.add_heading('Technical Skills Demonstrated', level=2)
    skills = [
        'Machine Learning: Model selection, training, hyperparameter tuning, evaluation',
        'Data Engineering: ETL pipelines, data validation, feature engineering',
        'Software Engineering: API design, error handling, code organization',
        'DevOps: Containerization, environment management, reproducibility',
        'Documentation: Professional README, API docs, technical writing'
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
    
    # ========== FOOTER ==========
    doc.add_page_break()
    doc.add_heading('Appendix: Repository Structure', level=1)
    doc.add_paragraph(
        'predictive_maintenance_project/\n'
        '├── src/\n'
        '│   ├── api.py (FastAPI server)\n'
        '│   ├── etl.py (Data pipeline)\n'
        '│   └── train.ipynb (Model training)\n'
        '├── models/\n'
        '│   └── rf_model.joblib (Trained model artifact)\n'
        '├── data/\n'
        '│   ├── ai4i2020_raw.csv (Raw dataset)\n'
        '│   └── processed_machine_data.csv (Processed data)\n'
        '├── docs/\n'
        '│   └── api-documentation.html (API reference)\n'
        '├── Dockerfile (Container configuration)\n'
        '├── environment.yml (Conda dependencies)\n'
        '├── requirements.txt (pip dependencies)\n'
        '└── README.md (Project documentation)'
    )
    
    doc.add_paragraph()
    # Metadata
    p = doc.add_paragraph()
    p.add_run('Document Generated: ').bold = True
    p.add_run(f'{datetime.now().strftime("%B %d, %Y")}')
    
    # Save document
    output_path = 'd:/GenAI Projects/predictive_maintenance_project/Predictive_Maintenance_Project_Reference.docx'
    doc.save(output_path)
    print(f"✅ Word document created successfully!")
    print(f"📄 Saved to: {output_path}")
    print(f"📊 Document includes: Executive Summary, Architecture, Results, Deployment, and Usage Guide")
    return output_path

if __name__ == '__main__':
    create_project_document()
